from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_resume(data):

    filename = "AI定制简历.docx"

    doc = Document()

    # ===== 标题 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run("个人简历")
    run.bold = True
    run.font.size = Pt(22)

    # ===== 基本信息 =====
    doc.add_heading("基本信息", level=2)

    basic_info = data.get("基本信息", {})

    doc.add_paragraph(
        f"姓名：{basic_info.get('姓名', '')}"
    )

    doc.add_paragraph(
        f"目标岗位：{basic_info.get('目标岗位', '')}"
    )

    # ===== 教育背景 =====
    doc.add_heading("教育背景", level=2)

    education = data.get("教育背景", "")

    if education:
        doc.add_paragraph(education)

    # ===== 专业技能 =====
    doc.add_heading("专业技能", level=2)

    skills = data.get("专业技能", [])

    for skill in skills:
        doc.add_paragraph(
            skill,
            style="List Bullet"
        )

    # ===== 项目经历 =====
    doc.add_heading("项目经历", level=2)

    projects = data.get("项目经历", [])

    for project in projects:
        doc.add_paragraph(
            project,
            style="List Bullet"
        )

    # ===== 个人优势 =====
    doc.add_heading("个人优势", level=2)

    advantages = data.get("个人优势", [])

    for advantage in advantages:
        doc.add_paragraph(
            advantage,
            style="List Bullet"
        )

    # ===== 保存 =====
    doc.save(filename)

    return {
        "success": True,
        "filename": filename
    }